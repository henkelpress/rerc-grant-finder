from __future__ import annotations

import csv
import hashlib
import json
import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Alignment, Font, PatternFill
from openpyxl.utils import get_column_letter
from openpyxl.worksheet.table import Table, TableStyleInfo


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "downloads"
DATE_TAG = date.today().isoformat()
PREFIX = "window.RERC_CATALOG = "
CASE_PREFIX = "window.RERC_CASE_STUDIES="
GREEN = "00573F"
DARK_GREEN = "173F35"
LEAF = "3E7C59"
GOLD = "F2C14E"
LIGHT_GREEN = "E9F2EC"
LIGHT_BLUE = "E6F1F5"


def load_js(path: Path, prefix: str) -> dict:
    raw = path.read_text(encoding="utf-8").strip()
    if not raw.startswith(prefix) or not raw.endswith(";"):
        raise ValueError(f"Unexpected JavaScript data format: {path.name}")
    return json.loads(raw[len(prefix) : -1])


def clean(value: object) -> str:
    return re.sub(r"\s+", " ", str(value or "")).strip()


def normalized_records() -> list[dict[str, str]]:
    catalog = load_js(ROOT / "data.js", PREFIX)
    cases = load_js(ROOT / "case_studies.js", CASE_PREFIX)
    records: list[dict[str, str]] = []
    for item in catalog["items"] + cases["items"]:
        item_type = clean(item.get("item_type"))
        records.append(
            {
                "Type": "Community Example" if item_type == "Case Study" else item_type,
                "Title": clean(item.get("title") or item.get("program")),
                "Organization": clean(item.get("organization") or item.get("agency")),
                "Status": clean(item.get("status")),
                "Availability or Year": clean(item.get("deadline_or_availability") or item.get("case_year")),
                "Geography": clean(item.get("case_state") or item.get("geography")),
                "Community": clean(item.get("case_place")),
                "Community Type": clean(item.get("case_place_type")),
                "Best For": clean(item.get("eligible_users")),
                "Project Stage": clean(item.get("project_stage")),
                "Topics": clean(item.get("topic_tags")),
                "Summary": clean(item.get("summary")),
                "Why It May Help": clean(item.get("why_it_matters")),
                "Amount or Cost": clean(item.get("amount_or_cost")),
                "Match or Cost": clean(item.get("match_or_cost")),
                "Official URL": clean(item.get("source_url")),
                "Last Checked": clean(item.get("last_checked")),
                "Record ID": clean(item.get("item_id")),
            }
        )
    return records


def add_hyperlink(paragraph, text: str, url: str) -> None:
    relationship = paragraph.part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), relationship)
    run = OxmlElement("w:r")
    properties = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "1B6A8F")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    properties.extend([color, underline])
    run.append(properties)
    content = OxmlElement("w:t")
    content.text = text
    run.append(content)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def set_repeat_table_header(row) -> None:
    row_properties = row._tr.get_or_add_trPr()
    repeat = OxmlElement("w:tblHeader")
    repeat.set(qn("w:val"), "true")
    row_properties.append(repeat)


def set_cell_shading(cell, color: str) -> None:
    properties = cell._tc.get_or_add_tcPr()
    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color)
    properties.append(shading)


def configure_document(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(9.5)
    for style_name, size, color in (
        ("Title", 30, GREEN),
        ("Heading 1", 20, GREEN),
        ("Heading 2", 12, DARK_GREEN),
    ):
        style = document.styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.color.rgb = RGBColor.from_string(color)
        style.font.bold = True
    for section in document.sections:
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        header = section.header.paragraphs[0]
        header.text = "Recreation Economy for Rural Communities"
        header.style = document.styles["Normal"]
        header.runs[0].font.color.rgb = RGBColor.from_string(GREEN)
        header.runs[0].font.bold = True
        footer = section.footer.paragraphs[0]
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer.add_run("RERC Community Explorer | Updated ")
        footer.add_run(DATE_TAG)


def add_record_table(document: Document, row: dict[str, str], index: int) -> None:
    title = document.add_heading(row["Title"], level=2)
    title.paragraph_format.space_before = Pt(8)
    title.paragraph_format.space_after = Pt(3)
    meta = document.add_paragraph()
    meta.paragraph_format.space_after = Pt(4)
    meta.add_run(row["Organization"]).bold = True
    details = [row["Status"], row["Geography"], row["Availability or Year"]]
    meta.add_run(" | " + " | ".join(item for item in details if item))

    table = document.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(1.35)
    table.columns[1].width = Inches(5.95)
    fields = [
        ("Summary", row["Summary"]),
        ("Best for", row["Best For"]),
        ("Why it may help", row["Why It May Help"]),
        ("Project stage", row["Project Stage"]),
        ("Topics", row["Topics"]),
        ("Amount or cost", row["Amount or Cost"]),
        ("Match or cost", row["Match or Cost"]),
        ("Official source", row["Official URL"]),
        ("Record ID", row["Record ID"]),
    ]
    for label, value in fields:
        if not value:
            continue
        cells = table.add_row().cells
        cells[0].text = label
        cells[0].paragraphs[0].runs[0].font.bold = True
        set_cell_shading(cells[0], LIGHT_GREEN if index % 2 else LIGHT_BLUE)
        if label == "Official source":
            add_hyperlink(cells[1].paragraphs[0], "Open official page", value)
            cells[1].paragraphs[0].add_run(f" ({value})")
        else:
            cells[1].text = value


def build_docx(records: list[dict[str, str]], path: Path) -> None:
    document = Document()
    configure_document(document)
    document.core_properties.title = "RERC Community Explorer: Funding, Resources, and Community Examples"
    document.core_properties.subject = "Public funding, technical resources, and official community examples"
    document.core_properties.author = "Recreation Economy for Rural Communities"
    document.core_properties.last_modified_by = "Recreation Economy for Rural Communities"
    document.core_properties.comments = ""
    title = document.add_paragraph(style="Title")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.add_run("RERC Community Explorer")
    subtitle = document.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_run = subtitle.add_run("Funding, Resources, and Community Examples")
    subtitle_run.bold = True
    subtitle_run.font.size = Pt(16)
    subtitle_run.font.color.rgb = RGBColor.from_string(DARK_GREEN)
    document.add_paragraph(
        "A public reference for rural communities developing outdoor recreation and community revitalization projects."
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    counts = {kind: sum(row["Type"] == kind for row in records) for kind in ("Funding", "Resource", "Community Example")}
    summary = document.add_table(rows=2, cols=3)
    summary.alignment = 1
    for index, kind in enumerate(("Funding", "Resource", "Community Example")):
        summary.cell(0, index).text = f"{counts[kind]:,}"
        summary.cell(1, index).text = {"Funding": "Funding opportunities", "Resource": "Resources", "Community Example": "Community examples"}[kind]
        for cell in (summary.cell(0, index), summary.cell(1, index)):
            set_cell_shading(cell, LIGHT_GREEN if index != 2 else LIGHT_BLUE)
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        summary.cell(0, index).paragraphs[0].runs[0].font.bold = True
        summary.cell(0, index).paragraphs[0].runs[0].font.size = Pt(18)

    note = document.add_paragraph()
    note.add_run("How to use this appendix. ").bold = True
    note.add_run(
        "Check current funding rules on the official page before applying. Community examples show approaches used elsewhere; they do not guarantee eligibility or results."
    )
    document.add_page_break()

    sections = (
        ("Funding Opportunities", "Funding"),
        ("Tools and Technical Resources", "Resource"),
        ("Community Examples", "Community Example"),
    )
    for section_index, (heading, kind) in enumerate(sections):
        if section_index:
            document.add_section(WD_SECTION.NEW_PAGE)
        document.add_heading(heading, level=1)
        matching = [row for row in records if row["Type"] == kind]
        intro = document.add_paragraph(f"{len(matching):,} public records. Use the filters in the online explorer to make a shorter community-specific appendix.")
        intro.runs[0].italic = True
        for index, row in enumerate(matching, 1):
            add_record_table(document, row, index)
    document.save(path)


def add_sheet(workbook: Workbook, title: str, rows: list[dict[str, str]]) -> None:
    sheet = workbook.create_sheet(title)
    headers = list(rows[0]) if rows else ["Type"]
    sheet.append(headers)
    for row in rows:
        sheet.append([row[header] for header in headers])
    for cell in sheet[1]:
        cell.fill = PatternFill("solid", fgColor=GREEN)
        cell.font = Font(color="FFFFFF", bold=True)
        cell.alignment = Alignment(wrap_text=True, vertical="center")
    sheet.freeze_panes = "A2"
    sheet.auto_filter.ref = sheet.dimensions
    sheet.sheet_view.showGridLines = False
    widths = {
        "A": 19, "B": 42, "C": 30, "D": 20, "E": 18, "F": 22, "G": 24, "H": 21,
        "I": 42, "J": 18, "K": 42, "L": 70, "M": 55, "N": 18, "O": 18, "P": 55, "Q": 15, "R": 24,
    }
    for column, width in widths.items():
        sheet.column_dimensions[column].width = width
    for row in sheet.iter_rows(min_row=2):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
        url_cell = row[headers.index("Official URL")]
        if url_cell.value:
            url_cell.hyperlink = url_cell.value
            url_cell.style = "Hyperlink"
    if rows:
        table = Table(displayName=re.sub(r"[^A-Za-z0-9]", "", title) + "Table", ref=sheet.dimensions)
        table.tableStyleInfo = TableStyleInfo(name="TableStyleMedium4", showRowStripes=True, showColumnStripes=False)
        sheet.add_table(table)


def build_xlsx(records: list[dict[str, str]], path: Path) -> None:
    workbook = Workbook()
    readme = workbook.active
    readme.title = "Read Me"
    readme.sheet_view.showGridLines = False
    readme["A1"] = "RERC Community Explorer"
    readme["A1"].font = Font(size=24, bold=True, color=GREEN)
    readme["A3"] = "Funding, Resources, and Community Examples"
    readme["A3"].font = Font(size=15, bold=True, color=DARK_GREEN)
    readme["A5"] = "Updated"
    readme["B5"] = DATE_TAG
    readme["A7"] = "Use"
    readme["B7"] = "Filter each sheet. Open the official URL to verify current details before applying or citing an example."
    readme["A9"] = "Public records"
    readme["B9"] = len(records)
    readme["A10"] = "Funding"
    readme["B10"] = sum(row["Type"] == "Funding" for row in records)
    readme["A11"] = "Resources"
    readme["B11"] = sum(row["Type"] == "Resource" for row in records)
    readme["A12"] = "Community examples"
    readme["B12"] = sum(row["Type"] == "Community Example" for row in records)
    readme.column_dimensions["A"].width = 24
    readme.column_dimensions["B"].width = 100
    for row in readme.iter_rows(min_row=1, max_row=12):
        for cell in row:
            cell.alignment = Alignment(wrap_text=True, vertical="top")
    add_sheet(workbook, "Funding", [row for row in records if row["Type"] == "Funding"])
    add_sheet(workbook, "Resources", [row for row in records if row["Type"] == "Resource"])
    add_sheet(workbook, "Community Examples", [row for row in records if row["Type"] == "Community Example"])
    workbook.properties.title = "RERC Community Explorer Master"
    workbook.properties.subject = "Funding, resources, and community examples"
    workbook.properties.creator = "Recreation Economy for Rural Communities"
    workbook.properties.lastModifiedBy = "Recreation Economy for Rural Communities"
    workbook.save(path)


def build_csv(records: list[dict[str, str]], path: Path) -> None:
    with path.open("w", encoding="utf-8-sig", newline="") as handle:
        writer = csv.DictWriter(handle, fieldnames=list(records[0]))
        writer.writeheader()
        writer.writerows(records)


def sha256(path: Path) -> str:
    return hashlib.sha256(path.read_bytes()).hexdigest()


def validate(records: list[dict[str, str]], docx: Path, xlsx: Path, csv_path: Path) -> dict:
    assert len(records) == 1197
    assert {row["Type"] for row in records} == {"Funding", "Resource", "Community Example"}
    assert all(row["Official URL"].startswith(("https://", "http://")) for row in records)
    serialized = json.dumps(records, ensure_ascii=False)
    assert not re.search(r"[A-Za-z]:\\|private_internal|needs_image_review|protos_case_id", serialized, re.I)
    assert not any(marker in serialized for marker in ("Ã", "Â", "â", "Æ", "ï¿½", "\ufffd"))
    workbook = load_workbook(xlsx, read_only=False, data_only=False)
    assert workbook.sheetnames == ["Read Me", "Funding", "Resources", "Community Examples"]
    assert workbook["Funding"].max_row == 660
    assert workbook["Resources"].max_row == 62
    assert workbook["Community Examples"].max_row == 478
    assert docx.stat().st_size > 100_000 and xlsx.stat().st_size > 100_000 and csv_path.stat().st_size > 100_000
    return {
        "status": "PASS",
        "updated": DATE_TAG,
        "records": len(records),
        "funding": 659,
        "resources": 61,
        "community_examples": 477,
        "source_sha256": {
            "data.js": sha256(ROOT / "data.js"),
            "case_studies.js": sha256(ROOT / "case_studies.js"),
        },
        "docx": {"file": docx.name, "bytes": docx.stat().st_size, "sha256": sha256(docx)},
        "xlsx": {"file": xlsx.name, "bytes": xlsx.stat().st_size, "sha256": sha256(xlsx)},
        "csv": {"file": csv_path.name, "bytes": csv_path.stat().st_size, "sha256": sha256(csv_path)},
    }


def main() -> int:
    OUTPUT.mkdir(parents=True, exist_ok=True)
    records = normalized_records()
    stem = f"RERC_Community_Explorer_Master_{DATE_TAG}"
    docx = OUTPUT / f"RERC_Community_Explorer_Appendix_{DATE_TAG}.docx"
    xlsx = OUTPUT / f"{stem}.xlsx"
    csv_path = OUTPUT / f"{stem}.csv"
    build_docx(records, docx)
    build_xlsx(records, xlsx)
    build_csv(records, csv_path)
    report = validate(records, docx, xlsx, csv_path)
    (OUTPUT / f"RERC_Community_Explorer_QA_{DATE_TAG}.json").write_text(
        json.dumps(report, indent=2) + "\n",
        encoding="utf-8",
    )
    print(json.dumps(report, indent=2))
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
